# plugins/docx_plugin.py
"""Плагин для анализа DOCX файлов"""

import os
from docx import Document

# Импортируем базовый класс
import sys
import os

sys.path.append(os.path.dirname(os.path.dirname(__file__)))
from core.plugin_base import DocumentPlugin


class DocxPlugin(DocumentPlugin):
    """Плагин для работы с DOCX файлами"""

    def __init__(self):
        super().__init__()
        self.name = "DOCX Анализатор"
        self.version = "1.0"
        self.supported_extensions = ['.docx', '.doc']

    def analyze(self, file_path):
        """Анализировать DOCX файл"""
        try:
            # Открываем документ
            doc = Document(file_path)

            # Собираем статистику
            stats = {
                'file_name': os.path.basename(file_path),
                'paragraphs': len(doc.paragraphs),
                'tables': len(doc.tables),
                'images': 0,  # Пока упрощенно
                'author': doc.core_properties.author or "Не указан",
                'created': str(doc.core_properties.created) or "Неизвестно"
            }

            # Извлекаем текст (первые 1000 символов)
            text_parts = []
            for para in doc.paragraphs[:20]:  # Первые 20 абзацев
                if para.text.strip():
                    text_parts.append(para.text)

            text_sample = "\n".join(text_parts)[:1000]

            return {
                "status": "success",
                "stats": stats,
                "text_sample": text_sample
            }

        except Exception as e:
            return {
                "status": "error",
                "message": f"Ошибка при анализе: {str(e)}"
            }