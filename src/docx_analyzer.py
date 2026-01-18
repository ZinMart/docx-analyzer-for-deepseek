"""
Модуль для анализа DOCX файлов
Содержит логику извлечения текста, изображений, таблиц и формул
"""

from docx import Document
import os

class DocxAnalyzer:
    """Анализатор DOCX файлов"""

    def __init__(self, file_path):
        self.file_path = file_path
        self.document = Document(file_path)
        self.stats = {
            'total_paragraphs': 0,
            'images': 0,
            'tables': 0,
            'formulas': 0,
            'other_elements': 0
        }

    def analyze(self):
        """Основной метод анализа файла"""
        # Подсчет параграфов
        self.stats['total_paragraphs'] = len(self.document.paragraphs)

        # Подсчет таблиц
        self.stats['tables'] = len(self.document.tables)

        # Подсчет изображений (базовый способ)
        # python-docx хранит изображения в rels документа
        try:
            rels = self.document.part.rels
            self.stats['images'] = sum(1 for rel in rels.values()
                                     if "image" in rel.target_ref)
        except:
            self.stats['images'] = 0

        return self.stats

    def extract_text(self):
        """Извлечение текста из документа"""
        full_text = []
        for paragraph in self.document.paragraphs:
            if paragraph.text.strip():  # Пропускаем пустые строки
                full_text.append(paragraph.text)

        return '\n'.join(full_text)

    def extract_images(self):
        """Извлечение изображений (заглушка)"""
        print("Метод extract_images() будет реализован позже")
        return []

    def extract_tables(self):
        """Извлечение таблиц (заглушка)"""
        print("Метод extract_tables() будет реализован позже")
        return []

    def extract_formulas(self):
        """Извлечение формул (заглушка)"""
        print("Метод extract_formulas() будет реализован позже")
        return []

    def get_basic_info(self):
        """Базовая информация о файле"""
        return {
            'filename': os.path.basename(self.file_path),
            'total_pages': 'N/A',  # python-docx не умеет считать страницы
            'author': self.document.core_properties.author or 'Не указан',
            'created': str(self.document.core_properties.created) or 'Не указано'
        }